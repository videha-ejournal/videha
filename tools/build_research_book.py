import json, html, re
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
DATA=json.loads((ROOT/'research/data/articles.json').read_text(encoding='utf-8'))['articles']
HTML=ROOT/'research/videha-scholar-research-book.html'; DOCX=ROOT/'research/Videha-Scholar-Research-Book.docx'
def auth(a): return ' / '.join(a.get('authors') or ['अज्ञात लेखक'])
def sortkey(a): return re.sub(r'^(डॉ\.?|डाॅ\.?|प्रो\.?|आचार्य|श्री|श्रीमती|पं\.?|पण्डित)\s*','',auth(a),flags=re.I).casefold()
def genre(a): return (a.get('classification') or 'अन्य शोध').split(';')[0].strip()
def pages(a): return f"{a.get('page_start')}–{a.get('page_end')}" if a.get('page_start') and a.get('page_end') else '—'
def date(a): return a.get('publication_date') or '—'
def issue(a): return 'अंक '+str(a.get('issue','—'))
def esc(s): return html.escape(str(s or ''),quote=True)
NARRATIVE = [
('विदेहक शोध-दृष्टि आ सम्पादकीय यात्रा', 'विदेह पहिल मैथिली पाक्षिक ई-जर्नलक रूपमे 2000 सँ भाषा, साहित्य आ समाजक मुक्त डिजिटल मंच बनबाक प्रयत्न करैत रहल अछि। सम्पादक गजेन्द्र ठाकुरक दृष्टिमे शोध केवल विश्वविद्यालयक परिधिमे सीमित कागजी विधा नहि, बल्कि मिथिलाक लोक-स्मृति, पञ्जी, गामक इतिहास, भाषा-विज्ञान, कला, नाटक, कानून, पर्यावरण आ समकालीन जीवनक प्रमाण-संग्रह सेहो अछि। अंक 1 सँ 447 धरि ई यात्रा रचना-प्रकाशनसँ बढ़ि कऽ एकटा दीर्घकालीन डिजिटल अभिलेख बनल अछि, जाहिमे मैथिली संग अंग्रेजी आ संस्कृतक विद्वत् लेख सेहो सुरक्षित अछि।'),
('संकलन-पद्धति आ समावेशक मानदण्ड', 'शोध-सूची बनबैत काल शीर्षकक चमक वा खण्डक नामपर निर्भर नहि कएल गेल अछि। प्रत्येक लेखक पूरा पाठ, लेखक-परिचय, अंक, तिथि, पृष्ठ, भाषा आ विषयगत कार्य देखल गेल। इतिहास, आलोचना, भाषा-विज्ञान, समाज-अध्ययन, मानवशास्त्र, लोक-संस्कृति, विधि, पर्यावरण, चिकित्सा-मानविकी, पञ्जी-अध्ययन आ दस्तावेजी सांस्कृतिक लेखनकेँ शोधयोग्य मानल गेल अछि। कथा, कविता, नाटक, सूचना, विज्ञप्ति, साधारण यात्रा-वृत्तान्त अथवा अल्प समाचारकेँ केवल विषय-साम्यक कारणेँ स्वतन्त्र शोध-लेख नहि बनाओल गेल। लम्बाइ एकटा सहायक संकेत अछि; निर्णायक कसौटी लेखक बौद्धिक कार्य, स्रोत-आधार आ विश्लेषणक गहराइ अछि।'),
('अंक-क्रममे विकास', 'आरम्भिक अंकमे भाषा, साहित्य, मिथिला-इतिहास आ समकालीन सामाजिक प्रश्नक आधारभूत लेख भेटैत अछि। अंक 11 सँ 100 धरि गजल, हाइकू, अनुवाद, नारी-विमर्श, बाल-साहित्य, नाटक, समीक्षा आ भाषिक प्रश्नपर विशेषांकक माध्यमे विषय-विस्तार भेल। अंक 101 सँ 230 धरि कानून, राज्य-आन्दोलन, वेब-पत्रकारिता, बीहनि कथा, लोक-संस्कृति आ प्रवासी अनुभवक लेख बढ़ल। अंक 231 सँ आगाँ गाम-अध्ययन, पर्यावरण, चिकित्सा-मानविकी, नव्य-न्याय, पञ्जी, अन्तरिक्ष-विज्ञान आ संस्थागत इतिहासक दस्तावेज जुड़ल। 37 विशेषांक एहि क्रमकेँ विषयगत गहनता देलक।'),
('भाषा-त्रयी : मैथिली, अंग्रेजी आ संस्कृत', 'मैथिली लेख सभसँ पैघ धारा अछि, मुदा अंग्रेजी लेख मिथिला-इतिहास, भाषिक नीति, चिकित्सा-मानविकी, प्रवास, जल-संकट आ तुलनात्मक साहित्यकेँ व्यापक पाठक धरि लऽ जाइत अछि। संस्कृत लेख व्याकरण, काव्यशास्त्र, न्याय, वेद, पाण्डुलिपि आ परम्परागत विद्याक आधुनिक पुनर्पाठ प्रस्तुत करैत अछि। भाषा-निर्धारण लेखमे प्रयुक्त मूल पाठक आधारपर कएल गेल अछि; अनुवादक भाषा आ मूल भाषाकेँ गड्ड-मड्ड नहि कएल गेल। एहि कारणेँ वर्तमान सत्यापित सूचीमे 777 मैथिली, 25 अंग्रेजी आ 8 संस्कृत लेख अछि।'),
('विद्वान, विषय आ विशेषांक', 'कैलाश कुमार मिश्रक गाम, चित्रकला आ मानवशास्त्रीय अध्ययन; आशीष अनचिन्हारक गजल आ वेब-पत्रकारिता; मुन्नाजीक बीहनि कथा; गजेन्द्र ठाकुरक आलोचना, भाषा-विज्ञान, नारी, दलित, क्वीयर, नाटक आ पद्य-विमर्श; विनीत उत्पलक संगोष्ठी-दस्तावेज; राधाकृष्ण चौधरीक अंग्रेजी मिथिला-इतिहास; प्रेमशंकर सिंह, योगेन्द्र प्रसाद यादव, बचेश्वर झा, रबीन्द्र नारायण मिश्र आ अनेक शोधकक लेख एहि भण्डारक बहुध्वनित स्वर अछि। हाइकू, गजल, बीहनि कथा, बाल-साहित्य, नाटक, समीक्षा, नारी, अनुवाद, वेब-पत्रकारिता, कला-विमर्श आ लेखक-विशेषांक सभ मिलि एकटा जीवित विषय-सूची बनबैत अछि।'),
('शोध-उपयोगिता आ सीमासभ', 'ई सूची शोधकर्ताकेँ प्राथमिक डिजिटल पाठ, स्थायी लिंक, अंक-परिचय, भाषा आ विधाक संकेत एक ठाम दैत अछि। तथापि ई विश्वविद्यालयीय peer-review सूची नहि; बहुत लेखक मूल स्रोत सीमित, OCR त्रुटिपूर्ण अथवा पृष्ठ-सूचना अपूर्ण भऽ सकैत अछि। किछु यात्रा-वर्णन, संस्मरण आ संस्थागत रिपोर्टक शोध-मूल्य सूक्ष्म पाठ-परीक्षणसँ बुझाइत अछि, तें चयनमे सम्पादकीय विवेक आवश्यक रहल। लेखक-नामक वर्तनी, उपाधि आ देवनागरी-अंग्रेजी रूपमे भिन्नता सेहो खोजमे चुनौती अछि।'),
('आगाँक विस्तार', 'आगामी चरणमे नूतन अंकक स्वचालित अभिलेखन, पूर्ण-पाठ OCR सुधार, लेखक-परिचयक प्रामाणिक प्राधिकरण, विषय-सम्बन्धक जाल, उद्धरण-निर्यात आ DOI-जकाँ स्थायी पहचान विकसित कएल जा सकैत अछि। पाठक अपन सुधार-सुझाव देथु, मुदा कोनो नव लेखकेँ शोध-सूचीमे जोड़बासँ पहिने मूल पाठ, लेखकत्व, अंक-स्रोत आ बौद्धिक स्वरूपक मानवीय जाँच अनिवार्य रहत। एहि तरह विदेहक डिजिटल विस्तार गति आ विश्वसनीयता दुनूकेँ संग राखत।'),
]
rows=sorted(DATA,key=lambda a:(sortkey(a),(a.get('title') or '').casefold())); groups=defaultdict(list)
for a in DATA: groups[genre(a)].append(a)
parts=['<!doctype html><html lang="mai"><head><meta charset="utf-8"><title>विदेह शोध-लेख : अंक १ सँ ४४७ धरि</title><style>@page{size:letter;margin:18mm}body{font-family:"Noto Serif Devanagari","Nirmala UI",serif;color:#18212b;line-height:1.5}h1{color:#163a5f}h2{color:#1f4d78;border-bottom:1px solid #b9c8d8}p{font-size:12px;text-align:justify}.note{background:#f1f5f8;border-left:4px solid #2e74b5;padding:10px}.book-tools{display:flex;flex-wrap:wrap;gap:.6rem;align-items:center;background:#faf6ee;border:1px solid #d8cfc0;padding:10px;margin:14px 0}.book-tools input{flex:1 1 18rem;padding:8px;font:inherit}.book-tools button{padding:7px 10px;font:inherit;cursor:pointer}.book-status{font-size:.9rem;color:#555}.videha-sr-only{position:absolute;width:1px;height:1px;margin:-1px;overflow:hidden;clip:rect(0,0,0,0)}table{border-collapse:collapse;width:100%;font-size:8.5px;page-break-inside:auto}thead{display:table-header-group}tr{page-break-inside:avoid}th,td{border:1px solid #b7c4d1;padding:4px;vertical-align:top}th{background:#e8eef5;color:#173b5f}a{color:#155a8a;text-decoration:none}</style></head><body>']
parts.append(f'<h1>विदेह शोध-लेख : अंक १ सँ ४४७ धरि</h1><p><b>लेखक आ सम्पादक:</b> Gajendra Thakur, Editor, Videha · ISSN 2229-547X · <a href="https://www.videha.co.in/">www.videha.co.in</a> · First Maithili Fortnightly eJournal</p><p>अन्तिम सत्यापित कॉर्पस : {len(DATA)} लेख · मैथिली {sum(a.get("language")=="mai" for a in DATA)} · अंग्रेजी {sum(a.get("language")=="en" for a in DATA)} · संस्कृत {sum(a.get("language")=="sa" for a in DATA)}</p>')
parts.append('<div class="note"><b>सम्पादकीय टिप्पणी।</b> एहि सूचीमे पूरा शोध, इतिहास, आलोचना, भाषा-विज्ञान, समाज-अध्ययन अथवा सांस्कृतिक अनुशीलनबला सामग्री राखल गेल अछि। कथा, कविता, सूचना आ मात्र अनुवादकेँ स्वतन्त्र शोध-लेख नहि गनल गेल अछि।</div>')
for heading, text in NARRATIVE:
 parts.append(f'<h2>{esc(heading)}</h2><p>{esc(text)}</p>')
parts.append('<div class="book-tools videha-a11y-bar videha-ai-standalone" role="region" aria-label="पृष्ठ सुनबाक, अनुवाद आ सहायक तकनीक नियंत्रण" data-pagefind-ignore="all"><button type="button" id="videha-tts-toggle" aria-pressed="false">🔊 सुनू · Listen</button><button type="button" id="videha-tts-stop" hidden>⏹ रोकू</button><span id="videha-tts-status" class="videha-sr-only" role="status" aria-live="polite"></span></div><div class="book-tools" role="search" aria-label="शोध लेख खोजू"><label for="book-search"><b>शोध लेख खोजू · Search</b></label><input id="book-search" type="search" placeholder="लेखक, शीर्षक, विधा, भाषा, अंक…" autocomplete="off"><span id="book-search-status" class="book-status" role="status" aria-live="polite">810 लेख</span></div>')
def hrow(n,a): return '<tr><td>%s</td><td>%s</td><td><a href="%s">%s</a></td><td>%s</td><td>%s</td><td>%s · %s</td><td>%s</td><td>%s</td></tr>'%(n,esc(auth(a)),esc(a.get('url')),esc(a.get('title')),esc(genre(a)),esc(a.get('language')),esc(issue(a)),esc(date(a)),esc(a.get('classification') or '—'),esc(pages(a)))
parts.append('<h2>अनुलग्नक १ : लेखकानुक्रमेण सम्पूर्ण सूची</h2><table><thead><tr><th>क्रम</th><th>लेखक</th><th>लेख</th><th>विधा</th><th>भाषा</th><th>अंक · तिथि</th><th>वर्गीकरण</th><th>पृष्ठ</th></tr></thead><tbody>')
for i,a in enumerate(rows,1): parts.append(hrow(i,a))
parts.append('</tbody></table><h2>अनुलग्नक २ : विधावार सम्पूर्ण सूची</h2>')
for g in sorted(groups,key=str.casefold):
 parts.append(f'<h3>{esc(g)} ({len(groups[g])})</h3><table><thead><tr><th>क्रम</th><th>लेखक</th><th>लेख</th><th>भाषा</th><th>अंक · तिथि</th><th>वर्गीकरण</th><th>पृष्ठ</th></tr></thead><tbody>')
 for i,a in enumerate(sorted(groups[g],key=lambda x:(sortkey(x),(x.get('title') or '').casefold())),1): parts.append('<tr><td>%d</td><td>%s</td><td><a href="%s">%s</a></td><td>%s</td><td>%s · %s</td><td>%s</td><td>%s</td></tr>'%(i,esc(auth(a)),esc(a.get('url')),esc(a.get('title')),esc(a.get('language')),esc(issue(a)),esc(date(a)),esc(a.get('classification') or '—'),esc(pages(a))))
 parts.append('</tbody></table>')
parts.append('<script src="../assets/js/videha-tts.js?v=20260818-hostfix2" defer></script><script src="../assets/js/videha-translate.js?v=20260827" defer data-videha-translate-standalone></script><script src="../assets/js/videha-access.js?v=20260827" defer></script><script>(function(){var q=document.getElementById("book-search"),st=document.getElementById("book-search-status"),rows=[].slice.call(document.querySelectorAll("table tbody tr"));function f(){var terms=(q.value||"").toLocaleLowerCase().trim().split(/\\s+/).filter(Boolean),n=0;rows.forEach(function(r){var ok=terms.every(function(t){return r.textContent.toLocaleLowerCase().indexOf(t)!==-1});r.hidden=!ok;if(ok)n++});st.textContent=(q.value?n+" मिलल · ":"")+"810 लेख"}q.addEventListener("input",f)})();</script>')
HTML.write_text(''.join(parts+['</body></html>']),encoding='utf-8')

def cell(c,text,bold=False,size=7):
 c.text='';p=c.paragraphs[0];p.paragraph_format.space_after=Pt(0);r=p.add_run(str(text));r.bold=bold;r.font.name='Noto Serif Devanagari';r.font.size=Pt(size)
def shade(c):
 sh=OxmlElement('w:shd');sh.set(qn('w:fill'),'E8EEF5');c._tc.get_or_add_tcPr().append(sh)
def table(doc,headers,data):
 t=doc.add_table(rows=1,cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.style='Table Grid';t.autofit=False
 for j,h in enumerate(headers): cell(t.rows[0].cells[j],h,True);shade(t.rows[0].cells[j])
 for row in data:
  cells=t.add_row().cells
  for j,v in enumerate(row): cell(cells[j],v)
doc=Document();s=doc.sections[0];s.page_width=Inches(8);s.page_height=Inches(11);s.top_margin=Inches(.75);s.bottom_margin=Inches(.75);s.left_margin=Inches(.9);s.right_margin=Inches(.75);s.gutter=Inches(.3)
doc.styles['Normal'].font.name='Noto Serif Devanagari';doc.styles['Normal'].font.size=Pt(10);doc.styles['Normal'].paragraph_format.space_after=Pt(5)
p=doc.add_paragraph();p.style='Title';p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run('विदेह शोध-लेख : अंक १ सँ ४४७ धरि')
doc.add_paragraph('लेखक आ सम्पादक: Gajendra Thakur, Editor, Videha · ISSN 2229-547X · www.videha.co.in · First Maithili Fortnightly eJournal').alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'अन्तिम सत्यापित कॉर्पस : {len(DATA)} लेख · मैथिली {sum(a.get("language")=="mai" for a in DATA)} · अंग्रेजी {sum(a.get("language")=="en" for a in DATA)} · संस्कृत {sum(a.get("language")=="sa" for a in DATA)}').alignment=WD_ALIGN_PARAGRAPH.CENTER
for i, (heading, text) in enumerate(NARRATIVE, 1):
 doc.add_heading(f'{i}. {heading}', 1); doc.add_paragraph(text)
doc.add_heading('अनुलग्नक १ : लेखकानुक्रमेण सम्पूर्ण सूची',1);table(doc,['क्रम','लेखक','लेख','विधा','भाषा','अंक · तिथि','वर्गीकरण','पृष्ठ'],[[i,auth(a),a.get('title'),genre(a),a.get('language'),f'{issue(a)} · {date(a)}',a.get('classification') or '—',pages(a)] for i,a in enumerate(rows,1)])
doc.add_heading('अनुलग्नक २ : विधावार सम्पूर्ण सूची',1)
for g in sorted(groups,key=str.casefold):
 doc.add_heading(f'{g} ({len(groups[g])})',2);table(doc,['क्रम','लेखक','लेख','भाषा','अंक · तिथि','वर्गीकरण','पृष्ठ'],[[i,auth(a),a.get('title'),a.get('language'),f'{issue(a)} · {date(a)}',a.get('classification') or '—',pages(a)] for i,a in enumerate(sorted(groups[g],key=lambda x:(sortkey(x),(x.get('title') or '').casefold())),1)])
doc.save(DOCX);print('created',HTML,DOCX,len(DATA))
