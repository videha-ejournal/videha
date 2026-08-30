from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import qrcode
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT=Path(__file__).resolve().parents[1]; R=ROOT/'research'; W,H=1600,2200
art=Image.open(R/'videha-research-book-cover-art.png').convert('RGB').resize((W,H),Image.Resampling.LANCZOS)
font=ImageFont.truetype('C:/Windows/Fonts/NIRMALA.TTF',48); bold=ImageFont.truetype('C:/Windows/Fonts/NIRMALAB.TTF',70)
def qr(url):
 q=qrcode.QRCode(version=4,box_size=8,border=3); q.add_data(url); q.make(fit=True); return q.make_image(fill_color='#111827',back_color='white').convert('RGB').resize((300,300),Image.Resampling.NEAREST)
def make_front():
 im=art.copy(); d=ImageDraw.Draw(im,'RGBA'); d.rounded_rectangle((170,500,1430,1450),radius=35,fill=(7,25,48,220),outline=(220,170,70,255),width=4); d.text((800,680),'विदेह शोध-सूची',font=bold,anchor='mm',fill='#f8e7ba'); d.text((800,820),'Videha Research Index',font=font,anchor='mm',fill='white'); d.text((800,980),'अंक १ सँ ४४७ धरि',font=bold,anchor='mm',fill='#f0b44c'); d.text((800,1140),'८१० सत्यापित शोध-लेख',font=font,anchor='mm',fill='white'); d.text((800,1300),'Gajendra Thakur / Editor, Videha',font=font,anchor='mm',fill='#f8e7ba'); return im
def make_back():
 im=art.copy(); d=ImageDraw.Draw(im,'RGBA'); d.rectangle((100,130,1500,600),fill=(7,25,48,225)); d.text((800,240),'विदेह शोध-संसार',font=bold,anchor='mm',fill='#f8e7ba'); d.text((800,370),'मिथिला, मैथिली, इतिहास आ संस्कृति',font=font,anchor='mm',fill='white'); im.paste(qr('https://www.videha.co.in/'),(240,760)); im.paste(qr('https://videha-ejournal.github.io/videha/'),(1060,760)); d=ImageDraw.Draw(im,'RGBA'); d.text((390,1110),'VIDEHA HOME',font=font,anchor='mm',fill='white'); d.text((1210,1110),'GITHUB HOME',font=font,anchor='mm',fill='white'); d.text((800,1350),'Scan to open the archive',font=font,anchor='mm',fill='#f8e7ba'); d.text((800,1850),'© Gajendra Thakur / Preeti Thakur',font=font,anchor='mm',fill='white'); d.text((800,1950),'Cover designed by Aum Gajendra Thakur',font=font,anchor='mm',fill='white'); return im
front=make_front(); back=make_back(); fp=R/'Videha-Scholar-Research-Book-Front-Cover.png'; bp=R/'Videha-Scholar-Research-Book-Back-Cover.png'; front.save(fp,dpi=(144,144)); back.save(bp,dpi=(144,144))
def pdf(img,path): c=canvas.Canvas(str(path),pagesize=(576,792)); c.drawImage(ImageReader(img),0,0,width=576,height=792); c.save()
pdf(front,R/'Videha-Scholar-Research-Book-Front-Cover.pdf'); pdf(back,R/'Videha-Scholar-Research-Book-Back-Cover.pdf')
src=R/'Videha-Scholar-Research-Book.docx'; doc=Document(src); first=doc.paragraphs[0]
def pic_before(path): p=first.insert_paragraph_before(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path),width=Inches(6.25),height=Inches(8.59)); p.paragraph_format.page_break_before=True
pic_before(fp); p=first.insert_paragraph_before(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('विदेह शोध-सूची\\nVideha Research Index\\n\\nअंक १ सँ ४४७ धरि\\n\\nGajendra Thakur, Editor, Videha').bold=True; p.paragraph_format.page_break_before=True
p=first.insert_paragraph_before(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('© Gajendra Thakur / Preeti Thakur\\n\\nCover designed by Aum Gajendra Thakur\\n\\nISSN 2229-547X\\nFirst Maithili Fortnightly eJournal').bold=True; p.paragraph_format.page_break_before=True
p=doc.add_paragraph(); p.paragraph_format.page_break_before=True; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(bp),width=Inches(6.25),height=Inches(8.59)); doc.save(R/'Videha-Scholar-Research-Book-with-Covers.docx')
print('created cover PNG/PDF and compiled DOCX')
